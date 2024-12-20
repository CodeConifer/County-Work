#Work Order Writer

#LIBRARIES
import openpyxl
import pandas as pd
from docx import Document
import os
from docx.shared import Pt
from docx.oxml.ns import qn

#SET FORMAT FOR THE PAGE
def format_paragraph(paragraph, font_name="Times New Roman", font_size=12, bold=True):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

#LOAD LOG, SET DF, LOAD TEMPLATE
log = r"C:\Users\nathan.stephens\OneDrive - Sussex County Government\Files\NKS_2024 P&D Log.xlsx"
df = pd.read_excel(log, sheet_name='W_R', header=4, dtype=str)
wr_temp = r"C:\Users\nathan.stephens\OneDrive - Sussex County Government\Files\Work Order_DOC TEMPLATE.docx"

#START COUNTERS AND LOGIC
created_files = []
created_count = 0

#LOAD LOG DATA
for index, row in df.iterrows():
    if row['STATUS'] == 'TO DO':
        # Load the Word template
        document = Document(wr_temp)

        #FORMAT SPECIFIC VALUES
        for paragraph in document.paragraphs:
            for key, value in row.items():
                placeholder = f'{{{{{key}}}}}'
                if placeholder in paragraph.text:
                    if key == 'Description':
                        descriptions = str(value).split(';')
                        formatted_description = "\n".join(descriptions)
                        paragraph.text = paragraph.text.replace(placeholder, formatted_description)

                    elif key == 'Type':
                        type_formatted = str(value).capitalize()
                        paragraph.text = paragraph.text.replace(placeholder, type_formatted)

                    elif key == 'Remarks':
                        remarks = str(value)
                        formatted_remarks = ''
                        entries = remarks.split("~")
                        for entry in entries:
                            parts = entry.split("][")
                            if len(parts) == 3:
                                action = parts[0].replace("[", "").replace("]", "").strip()
                                size = parts[1].replace("[", "").replace("]", "").strip()
                                description = parts[2].replace("[", "").replace("]", "").strip()
                                parts = description.split(';')
                                formatted_remarks += (
                                    f"{action.capitalize()} Parcel: {action.split(';')[-1]}\n"
                                    f"Size: {size} Ac.\n"
                                    f"Description: {parts[0]}\n" +
                                    (f"{parts[1]}\n" if len(parts) > 1 else "") +
                                    (f"{parts[2]}\n" if len(parts) > 2 else "") +
                                    (f"{parts[3]}\n" if len(parts) > 3 else "") +
                                    "\n"
                                )
                            else:
                                formatted_remarks += f"{entry.strip()}\n\n"
                        paragraph.text = paragraph.text.replace(placeholder, formatted_remarks)

                    else:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
                    
                    #APPLY FORMATTING TO WR
                    format_paragraph(paragraph)

        #ADD DATA TO FOOTER
        section = document.sections[0]
        footer = section.footer
        for paragraph in footer.paragraphs:
            if "{{Date Rec}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{Date Rec}}", str(row['Date Rec']))

        #SAVE WR
        wr_final = f"{row['ID']}.docx"
        output_path = r"C:\Users\nathan.stephens\OneDrive - Sussex County Government\Files\Work Requests"
        output_file = os.path.join(output_path, wr_final)
        document.save(output_file)

        #TOTAL CREATED FILES
        created_files.append(wr_final)
        created_count += 1

#PRINT RESULTS
if created_count > 0:
    print(f"\nWork Orders created: {created_count}")
    print("Created:")
    for file in created_files:
        print(file)